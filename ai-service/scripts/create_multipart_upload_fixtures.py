"""Create deterministic long binary fixtures for /internal/process multi-upload tests."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate


ROOT = Path(__file__).resolve().parent.parent / "test-fixtures" / "multipart_upload"

DOCX_REQUIREMENTS = [
    "The customer workspace shall let an account owner create a project and invite named collaborators with a project-scoped role.",
    "The service shall issue a single-use email invitation link that expires after twenty-four hours and records its redemption time.",
    "The workspace shall require multi-factor authentication for administrators before they can change organization settings or billing contacts.",
    "The system shall record immutable audit events for invitation creation, role changes, sign-in failures, and export requests.",
    "The audit search screen shall filter by actor, action, target project, and a caller-selected date range.",
    "The export service shall produce CSV and PDF audit reports and include the applied filters in each generated artifact.",
    "The application shall retain exported reports for thirty days and allow only administrators to retrieve a retained report.",
    "The notification service shall alert account owners when a new administrator role is granted or when an export is downloaded.",
]

UNRELATED_GARDENING = [
    "UNRELATED MIDDLE SECTION - Community garden meeting minutes: volunteers discussed tomato seedlings, compost delivery, and a Saturday watering schedule.",
    "The gardening group selected rosemary, basil, and marigolds for the east planter and agreed to label each bed with a weatherproof sign.",
    "The meeting also recorded a recipe exchange for lemon herb dressing and a reminder to bring gloves to the next cleanup day.",
]

PDF_REQUIREMENTS = [
    "The operations portal shall show a queue of pending support cases and allow an assigned analyst to change case status with a reason.",
    "Each status transition shall preserve the prior value, the acting user, the timestamp, and a human-readable rationale in case history.",
    "The portal shall enforce a four-hour response target for high-priority cases and display a warning when less than one hour remains.",
    "Supervisors shall configure escalation rules by customer tier, business hours calendar, and unresolved case age.",
    "The system shall send escalation notifications to the primary on-call group and retain delivery outcomes for troubleshooting.",
    "Analysts shall attach sanitized diagnostic files to a case; attachments shall be virus-scanned before they become available to other users.",
    "The reporting view shall summarize response-time compliance by team, priority, and calendar month without exposing customer credentials.",
    "Administrators shall export a monthly operations report and the report shall identify its source period and generation timestamp.",
]

UNRELATED_TRAVEL = [
    "UNRELATED MIDDLE SECTION - Weekend travel notes: the writer compared museum opening hours, train schedules, and a list of quiet cafes near the river.",
    "A separate paragraph describes packing a rain jacket, choosing postcards, and taking photographs of bridges at sunset.",
    "These travel notes are intentionally unrelated to software requirements and should remain attributable to this PDF source.",
]


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_docx_section(document: Document, heading: str, paragraphs: list[str], repeats: int) -> None:
    document.add_heading(heading, level=1)
    for cycle in range(repeats):
        for index, text in enumerate(paragraphs, start=1):
            body = f"{text} Review note {cycle + 1}.{index}: this requirement is included for multipart extraction coverage."
            document.add_paragraph(body)


def create_docx(path: Path) -> None:
    document = Document()
    _configure_docx(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Customer Workspace Requirements Specification")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)
    document.add_paragraph("Fixture: multipart document upload regression. Version 1.0.")
    _add_docx_section(document, "Workspace and access requirements", DOCX_REQUIREMENTS[:4], repeats=3)
    _add_docx_section(document, "Unrelated material", UNRELATED_GARDENING, repeats=2)
    _add_docx_section(document, "Audit and export requirements", DOCX_REQUIREMENTS[4:], repeats=3)
    document.add_heading("Acceptance review", level=1)
    for item in DOCX_REQUIREMENTS:
        document.add_paragraph(item, style="List Bullet")
    document.save(path)


def create_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("FixtureTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#0B2545"), spaceAfter=16)
    heading = ParagraphStyle("FixtureHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=8)
    body = ParagraphStyle("FixtureBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, spaceAfter=8)
    document = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    story = [
        Paragraph("Operations Case Management Requirements", title),
        Paragraph("Fixture: multipart document upload regression. Version 1.0.", body),
        Paragraph("Case workflow requirements", heading),
    ]
    for cycle in range(3):
        for index, text in enumerate(PDF_REQUIREMENTS[:4], start=1):
            story.append(Paragraph(f"{text} Validation note {cycle + 1}.{index}: this requirement is repeated to create a realistic long source.", body))
    story.append(Paragraph("Unrelated material", heading))
    for cycle in range(3):
        for text in UNRELATED_TRAVEL:
            story.append(Paragraph(f"{text} Travel note iteration {cycle + 1}.", body))
    story.append(Paragraph("Escalation and reporting requirements", heading))
    for cycle in range(3):
        for index, text in enumerate(PDF_REQUIREMENTS[4:], start=1):
            story.append(Paragraph(f"{text} Review note {cycle + 1}.{index}: this requirement supports long-document chunking validation.", body))
    document.build(story)


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    create_docx(ROOT / "customer_workspace_requirements.docx")
    create_pdf(ROOT / "operations_case_management.pdf")


if __name__ == "__main__":
    main()
