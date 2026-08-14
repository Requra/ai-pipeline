import io
import pytest
import docx
from app.nodes.ingest import _extract_docx, ingest_node
from app.nodes.parse_to_chunks import parse_to_chunks_node, _chunk_docx_paragraphs
from app.schemas.items import SourceChunk


def test_docx_extract_and_chunking_fallback():
    doc = docx.Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("This is paragraph one.")
    doc.add_paragraph("This is paragraph two.")
    
    stream = io.BytesIO()
    doc.save(stream)
    raw_bytes = stream.getvalue()
    
    res = _extract_docx(raw_bytes)
    # Check that it either extracted paragraphs or converted to PDF (on system with Word/LibreOffice)
    assert len(res) == 3


@pytest.mark.asyncio
async def test_ingest_and_parse_to_chunks_docx_fallback_e2e(base_state, monkeypatch):
    async def fake_relevance(_: str):
        from app.nodes.ingest import RelevanceCheck
        return RelevanceCheck(
            is_useful=True,
            relevance_score=0.9,
            reason="Software specification",
        )
    monkeypatch.setattr("app.nodes.ingest._run_relevance_check", fake_relevance)
    # DOCX rendering is deliberately deferred for MVP and must not run.
    def conversion_must_not_run(_):
        raise AssertionError("DOCX-to-PDF conversion must remain disabled for MVP")

    monkeypatch.setattr(
        "app.nodes.ingest.convert_docx_to_pdf",
        conversion_must_not_run,
    )

    doc = docx.Document()
    doc.add_heading("E2E Section", level=1)
    doc.add_paragraph("E2E paragraph content functional requirements system.")
    
    stream = io.BytesIO()
    doc.save(stream)
    raw_bytes = stream.getvalue()
    
    state = base_state.copy()
    state["raw_bytes"] = raw_bytes
    state["file_type"] = "docx"
    state["metadata"] = {"filename": "test.docx"}
    
    ingest_res = await ingest_node(state)
    assert ingest_res["status"] == "ready_for_chunking"
    assert len(ingest_res["source_documents"]) == 1
    doc_meta = ingest_res["source_documents"][0]
    assert doc_meta["docx_paragraphs"] is not None
    
    state.update(ingest_res)
    chunk_res = await parse_to_chunks_node(state)
    assert chunk_res["status"] == "chunks_parsed"
    assert len(chunk_res["chunks"]) > 0
    c = chunk_res["chunks"][0]
    assert c.page_number is None
    assert c.heading == "E2E Section"
    assert c.section == "E2E Section"
