"""
Purpose-designed multi-source E2E fixture generator.

Generates heterogeneous sources (PDF, DOCX, TXT, WAV) where each source
contains a unique, non-overlapping semantic requirement anchor and negative
control prose.

Fixture Definitions:
  - ALPHA (PDF):
      Requirement: "ALPHA-LOCK-731: The platform shall suspend authentication for exactly 17 minutes after seven consecutive invalid password attempts."
      Anchor: "7 failed attempts -> 17-minute authentication suspension"
      Negative control: "The team ordered pepperoni pizza after the security workshop."

  - BETA (DOCX):
      Requirement: "BETA-INVOICE-842: After a payment is confirmed, the billing system shall generate a downloadable tax invoice containing the transaction reference and VAT amount."
      Anchor: "confirmed payment -> downloadable VAT invoice"
      Negative control: "The finance department meets every Tuesday at 10 AM in room 302."

  - GAMMA (TXT):
      Requirement: "GAMMA-AUDIT-953: Administrators shall export an audit CSV containing the actor, action timestamp, previous value and new value."
      Anchor: "admin audit CSV -> actor + timestamp + previous/new values"
      Negative control: "The server room air conditioner was serviced last Thursday."

  - DELTA (Audio WAV):
      Requirement: "DELTA-VOICE-614: Warehouse supervisors need a warning when refrigerated storage temperature stays above eight degrees Celsius for five minutes."
      Anchor: "temperature > 8°C for 5 minutes -> warehouse supervisor warning"
      Negative control: "The warehouse inventory count was scheduled for Friday afternoon."
"""

from __future__ import annotations

import io
import math
import struct
import wave
from typing import Dict, Any, List

import docx
import fitz  # PyMuPDF


ALPHA_TEXT = (
    "ALPHA-LOCK-731\n\n"
    "The platform shall suspend authentication for exactly 17 minutes after seven consecutive invalid password attempts.\n\n"
    "The team ordered pepperoni pizza after the security workshop."
)

BETA_TEXT = (
    "BETA-INVOICE-842\n\n"
    "After a payment is confirmed, the billing system shall generate a downloadable tax invoice containing the transaction reference and VAT amount.\n\n"
    "The finance department meets every Tuesday at 10 AM in room 302."
)

GAMMA_TEXT = (
    "GAMMA-AUDIT-953\n\n"
    "Administrators shall export an audit CSV containing the actor, action timestamp, previous value and new value.\n\n"
    "The server room air conditioner was serviced last Thursday."
)

DELTA_TRANSCRIPT = (
    "DELTA-VOICE-614. "
    "Warehouse supervisors need a warning when refrigerated storage temperature stays above eight degrees Celsius for five minutes. "
    "The warehouse inventory count was scheduled for Friday afternoon."
)


def get_alpha_pdf_bytes() -> bytes:
    """Generate a valid PDF containing the ALPHA security requirement."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    page.insert_text(
        (50, 72),
        "Requra System Security Specification\n\n" + ALPHA_TEXT,
        fontsize=11,
    )
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def get_beta_docx_bytes() -> bytes:
    """Generate a valid DOCX containing the BETA billing requirement."""
    doc = docx.Document()
    doc.add_heading("Billing System Requirements", level=1)
    for paragraph in BETA_TEXT.split("\n\n"):
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def get_gamma_txt_bytes() -> bytes:
    """Generate a valid UTF-8 text file containing the GAMMA audit requirement."""
    return ("Requra System Audit Specification\n\n" + GAMMA_TEXT).encode("utf-8")


def get_delta_wav_bytes(duration_seconds: float = 1.0, sample_rate: int = 16000) -> bytes:
    """Generate a valid 16-bit mono PCM WAV audio buffer."""
    buf = io.BytesIO()
    with wave.open(buf, "wb") as wav:
        wav.setnchannels(1)
        wav.setsampwidth(2)
        wav.setframerate(sample_rate)
        total_frames = int(duration_seconds * sample_rate)
        freq = 440.0  # 440 Hz standard tone
        frames = bytearray()
        for i in range(total_frames):
            val = int(32767.0 * 0.5 * math.sin(2.0 * math.pi * freq * (i / sample_rate)))
            frames.extend(struct.pack("<h", val))
        wav.writeframes(frames)
    return buf.getvalue()


def get_irrelevant_pdf_bytes() -> bytes:
    """Generate a valid PDF with irrelevant non-software content."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text(
        (50, 72),
        "Traditional French Croissant Recipe\n\n"
        "Ingredients:\n"
        "- 500g all-purpose flour\n"
        "- 10g salt\n"
        "- 50g sugar\n"
        "- 300g unsalted butter for laminating\n\n"
        "Instructions: Mix dry ingredients with water and knead into a smooth dough. "
        "Chill overnight before folding with butter layers.",
        fontsize=11,
    )
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def get_corrupted_bytes() -> bytes:
    """Generate corrupted, unparseable binary content."""
    return b"%PDF-1.4\n\x00\xff\xfe\x00INVALID_CORRUPTED_BINARY_STREAM_%%EOF"


def get_all_four_sources_manifest() -> List[Dict[str, Any]]:
    """Return manifest list for all four canonical test sources."""
    return [
        {
            "document_id": "doc_alpha_pdf",
            "filename": "source-alpha.pdf",
            "file_type": "document",
            "mime_type": "application/pdf",
            "raw_bytes": get_alpha_pdf_bytes(),
            "expected_anchor": "17 minutes",
            "semantic_key": "ALPHA",
        },
        {
            "document_id": "doc_beta_docx",
            "filename": "source-beta.docx",
            "file_type": "document",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "raw_bytes": get_beta_docx_bytes(),
            "expected_anchor": "VAT",
            "semantic_key": "BETA",
        },
        {
            "document_id": "doc_gamma_txt",
            "filename": "source-gamma.txt",
            "file_type": "document",
            "mime_type": "text/plain",
            "raw_bytes": get_gamma_txt_bytes(),
            "expected_anchor": "audit CSV",
            "semantic_key": "GAMMA",
        },
        {
            "document_id": "doc_delta_audio",
            "filename": "source-delta.wav",
            "file_type": "audio",
            "mime_type": "audio/wav",
            "raw_bytes": get_delta_wav_bytes(),
            "expected_anchor": "eight degrees",
            "semantic_key": "DELTA",
        },
    ]
