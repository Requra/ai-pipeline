"""
Reusable document extraction, text normalization, PII redaction, and relevance check services.
Decoupled from LangGraph nodes to allow safe reuse across API, worker, and background pipelines.
"""

from __future__ import annotations

import io
import json
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import docx
from pydantic import BaseModel, Field

from app.config import settings
from app.llm import get_llm
from app.prompts.loader import load_prompt
from app.prompts.registry import PromptId

logger = logging.getLogger(__name__)

MIN_TEXT_LENGTH = 50
RELEVANCE_SNIPPET_CHARS = 2000

EMAIL_PATTERN = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_PATTERN = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)")
OPENAI_KEY_PATTERN = re.compile(r"\b(?:sk-proj-[a-zA-Z0-9_]{32,}|sk-[a-zA-Z0-9]{32,})\b")
AWS_KEY_PATTERN = re.compile(r"\b(?:AKIA|ASCA)[A-Z0-9]{16}\b")
GITHUB_KEY_PATTERN = re.compile(r"\b(?:ghp|gho|ghu|ghs|ghr)_[a-zA-Z0-9]{36}\b")
HF_KEY_PATTERN = re.compile(r"\bhf_[a-zA-Z0-9]{34}\b")
GOOGLE_KEY_PATTERN = re.compile(r"\bAIzaSy[a-zA-Z0-9_\-]{33}\b")
GENERIC_SECRET_PATTERN = re.compile(
    r"(?i)\b(api_key|secret_key|private_key|access_token|db_password)(\s*=\s*['\"]?)([a-zA-Z0-9_\-]{16,})(['\"]?)\b"
)
CREDIT_CARD_CANDIDATE_PATTERN = re.compile(r"\b(?:\d[ -]*?){13,16}\b")


class RelevanceCheckResult(BaseModel):
    decision: str = Field(
        default="relevant",
        description="Tri-state relevance decision: 'relevant', 'uncertain', or 'irrelevant'."
    )
    is_useful: bool = Field(
        default=True,
        description="True if the source contains potential requirements or if relevance is uncertain; False only when high-confidence irrelevant."
    )
    confidence: float = Field(
        default=0.5,
        description="Confidence in this classification between 0.0 and 1.0."
    )
    relevance_score: float = Field(
        default=0.5,
        description="Relevance score between 0.0 and 1.0."
    )
    reason: str = Field(
        default="",
        description="Short reason explaining acceptance, uncertainty, or rejection."
    )
    evidence: list[str] = Field(
        default_factory=list,
        description="Excerpts from source supporting the decision."
    )
    signals: Optional[dict[str, bool]] = Field(
        default=None,
        description="Detected requirements signals (requirements, business_rules, constraints, workflows, decisions, stakeholders)."
    )
    method: str = Field(
        default="llm",
        description="Classification method used: 'llm', 'deterministic_heuristic', or 'fail_open_fallback'."
    )


def _normalize_text(text: str) -> str:
    """Normalize whitespace for stable downstream parsing and idempotent outputs."""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    lines: list[str] = []
    previous_was_blank = False

    for line in normalized.split("\n"):
        compact = re.sub(r"[ \t]+", " ", line).strip()
        if not compact:
            if not previous_was_blank:
                lines.append("")
            previous_was_blank = True
            continue

        lines.append(compact)
        previous_was_blank = False

    return "\n".join(lines).strip()


def _is_luhn_valid(number: str) -> bool:
    digits = [int(c) for c in number if c.isdigit()]
    if not digits:
        return False
    checksum = 0
    reverse_digits = digits[::-1]
    for i, d in enumerate(reverse_digits):
        if i % 2 == 1:
            doubled = d * 2
            checksum += doubled - 9 if doubled > 9 else doubled
        else:
            checksum += d
    return checksum % 10 == 0


def _mask_pii(text: str) -> tuple[str, dict[str, int]]:
    stats: dict[str, int] = {}

    def _replace_card(match: re.Match) -> str:
        val = match.group(0)
        raw_digits = re.sub(r"\D", "", val)
        if 13 <= len(raw_digits) <= 16 and _is_luhn_valid(raw_digits):
            stats["CREDIT_CARD"] = stats.get("CREDIT_CARD", 0) + 1
            return "[REDACTED_CREDIT_CARD]"
        return val

    masked = CREDIT_CARD_CANDIDATE_PATTERN.sub(_replace_card, text)

    def _replace_secret(match: re.Match) -> str:
        key_name = match.group(1)
        stats["API_SECRET"] = stats.get("API_SECRET", 0) + 1
        return f"{key_name}=[REDACTED_SECRET]"

    masked = GENERIC_SECRET_PATTERN.sub(_replace_secret, masked)

    def _replace_known_token(match: re.Match, token_name: str) -> str:
        stats[token_name] = stats.get(token_name, 0) + 1
        return f"[REDACTED_{token_name}]"

    masked = OPENAI_KEY_PATTERN.sub(lambda m: _replace_known_token(m, "OPENAI_KEY"), masked)
    masked = AWS_KEY_PATTERN.sub(lambda m: _replace_known_token(m, "AWS_KEY"), masked)
    masked = GITHUB_KEY_PATTERN.sub(lambda m: _replace_known_token(m, "GITHUB_TOKEN"), masked)
    masked = HF_KEY_PATTERN.sub(lambda m: _replace_known_token(m, "HF_TOKEN"), masked)
    masked = GOOGLE_KEY_PATTERN.sub(lambda m: _replace_known_token(m, "GOOGLE_API_KEY"), masked)

    email_matches = EMAIL_PATTERN.findall(masked)
    if email_matches:
        stats["EMAIL"] = len(email_matches)
        masked = EMAIL_PATTERN.sub("[EMAIL]", masked)

    phone_matches = PHONE_PATTERN.findall(masked)
    if phone_matches:
        stats["PHONE"] = len(phone_matches)
        masked = PHONE_PATTERN.sub("[PHONE]", masked)

    return masked, stats


def _extract_pdf(raw_bytes: bytes) -> tuple[str, Optional[str]]:
    if not raw_bytes:
        return "", "INGEST_FAILED: missing PDF bytes"
    try:
        pages: list[str] = []
        with fitz.open(stream=raw_bytes, filetype="pdf") as document:
            for page in document:
                pages.append(page.get_text())
        return "\f".join(pages), None
    except Exception as e:
        return "", f"INGEST_FAILED: PDF extraction error ({e})"


def _extract_docx(raw_bytes: bytes) -> tuple[str, Optional[str], Optional[list[dict]]]:
    if not raw_bytes:
        return "", "INGEST_FAILED: missing DOCX bytes", None
    try:
        doc = docx.Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n\n".join(paragraphs)
        paragraphs_data = [{"text": p, "paragraph_index": i} for i, p in enumerate(paragraphs)]
        return _normalize_text(text), None, paragraphs_data
    except Exception as e:
        return "", f"INGEST_FAILED: DOCX extraction error ({e})", None


RELEVANCE_MAX_TOTAL_CHARS = 3000


def _sample_representative_text(text: str, max_chars: int = RELEVANCE_MAX_TOTAL_CHARS) -> str:
    """Sample head, middle, and tail spans of text so requirements situated after

    introductions, greetings, or agenda setup are not missed.
    """
    clean_text = text.replace("\f", " ").strip()
    if len(clean_text) <= max_chars:
        return clean_text

    window = max_chars // 3
    head = clean_text[:window].strip()
    mid_start = max(0, (len(clean_text) // 2) - (window // 2))
    mid = clean_text[mid_start : mid_start + window].strip()
    tail = clean_text[-window:].strip()

    return f"{head}\n\n[... middle section ...]\n\n{mid}\n\n[... ending section ...]\n\n{tail}"


def _conservative_deterministic_relevance(
    snippet: str,
    *,
    reason_prefix: str = "Deterministic analysis"
) -> RelevanceCheckResult:
    """Conservative, domain-agnostic deterministic relevance check.

    Analyzes behavioral patterns, conditions, business rules, operational constraints,
    and stakeholder needs across all domains (agriculture, healthcare, retail, IoT, etc.).
    Fails open to 'uncertain' whenever ambiguous or weak rather than falsely rejecting.
    """
    lowered = snippet.lower().strip()

    if not lowered or len(lowered) < MIN_TEXT_LENGTH:
        return RelevanceCheckResult(
            decision="uncertain" if len(lowered) >= 20 else "irrelevant",
            is_useful=len(lowered) >= 20,
            confidence=0.5 if len(lowered) >= 20 else 0.85,
            relevance_score=0.3 if len(lowered) >= 20 else 0.0,
            reason=f"{reason_prefix}: source content is empty or near-empty ({len(lowered)} chars).",
            evidence=[],
            signals={
                "requirements": False,
                "business_rules": False,
                "constraints": False,
                "workflows": False,
                "decisions": False,
                "stakeholders": False,
            },
            method="deterministic_heuristic",
        )

    # 1. Structural / Behavioral signals
    # Requirement modal / obligation verbs
    modal_matches = len(re.findall(r"\b(must|shall|should|need to|needs to|required to|has to|have to|cannot|can not|will be able to|able to|must not|should not)\b", lowered))

    # Conditional / trigger patterns
    condition_matches = len(re.findall(r"\b(when|if|whenever|unless|before|after|upon|in case|once|as soon as|where)\b", lowered))

    # Action & operation verbs
    action_matches = len(re.findall(r"\b(open|close|start|stop|begin|allow|prevent|trigger|notify|alert|send|receive|verify|calculate|display|show|track|record|override|approve|reject|confirm|generate|log|update|create|delete|manage|monitor|control|disable|enable|restrict)\b", lowered))

    # Quantitative / threshold / constraint patterns
    constraint_matches = len(re.findall(r"\b(below|above|exceed|exceeds|at least|at most|maximum|minimum|limit|within|seconds|minutes|hours|days|percent|%|threshold|range|tolerance|meters|liters|gallons|degrees)\b", lowered))

    # Domain roles / actors
    actor_matches = len(re.findall(r"\b(user|users|admin|manager|operator|customer|client|patient|nurse|doctor|driver|dispatcher|farmer|technician|worker|agent|officer|coach|athlete|system|device|sensor|valve|service|portal|app)\b", lowered))

    # Explicit artifact indicators (supporting signal)
    artifact_matches = len(re.findall(r"\b(requirement|spec|specification|workflow|process|rule|policy|criteria|feature|story|task|backlog|epic|module|database|api|integration|interface)\b", lowered))

    # Obvious non-project garbage indicators (recipes, songs/chords, lorem ipsum)
    recipe_matches = len(re.findall(r"\b(tablespoon|tablespoons|teaspoon|teaspoons|cups of|flour|sugar|baking powder|preheat|oven to|simmer|stir well|bake for|degrees fahrenheit|recipe)\b", lowered))
    lyrics_matches = len(re.findall(r"\b(chorus|verse 1|verse 2|intro:|outro:|la la la|oh yeah|repeat chorus)\b", lowered))
    lorem_matches = len(re.findall(r"\b(lorem ipsum|dolor sit amet|consectetur adipiscing|vestibulum|pellentesque)\b", lowered))

    has_requirements = (modal_matches >= 1 and (action_matches >= 1 or actor_matches >= 1)) or artifact_matches >= 1
    has_business_rules = (condition_matches >= 1 and (modal_matches >= 1 or action_matches >= 1)) or (constraint_matches >= 1 and (modal_matches >= 1 or actor_matches >= 1))
    has_constraints = constraint_matches >= 1 and (modal_matches >= 1 or actor_matches >= 1 or condition_matches >= 1)
    has_workflows = (condition_matches >= 1 and action_matches >= 1) or (action_matches >= 2 and actor_matches >= 1)
    has_decisions = bool(re.search(r"\b(decided|agreed|chosen|selected|decision|approved|rejected)\b", lowered))
    has_stakeholders = actor_matches >= 1

    signals = {
        "requirements": bool(has_requirements),
        "business_rules": bool(has_business_rules),
        "constraints": bool(has_constraints),
        "workflows": bool(has_workflows),
        "decisions": bool(has_decisions),
        "stakeholders": bool(has_stakeholders),
    }

    positive_score = (
        (1.5 if has_requirements else 0) +
        (1.0 if has_business_rules else 0) +
        (0.8 if has_workflows else 0) +
        (0.7 if has_constraints else 0) +
        (0.5 if has_stakeholders else 0) +
        (0.5 if artifact_matches >= 1 else 0)
    )

    # Definite garbage detection: strong non-project markers with zero requirements
    is_definite_garbage = (
        (recipe_matches >= 2 and not has_requirements and modal_matches == 0) or
        (lyrics_matches >= 2 and not has_requirements and modal_matches == 0) or
        (lorem_matches >= 2 and not has_requirements)
    )

    if is_definite_garbage:
        return RelevanceCheckResult(
            decision="irrelevant",
            is_useful=False,
            confidence=0.9,
            relevance_score=0.1,
            reason=f"{reason_prefix}: content identified as non-project material (recipe/lyrics/lorem).",
            evidence=[],
            signals=signals,
            method="deterministic_heuristic",
        )

    if positive_score >= 1.5:
        return RelevanceCheckResult(
            decision="relevant",
            is_useful=True,
            confidence=min(0.95, 0.65 + (positive_score * 0.08)),
            relevance_score=min(1.0, 0.75 + (positive_score * 0.05)),
            reason=f"{reason_prefix}: identified requirements and domain operational rules.",
            evidence=[],
            signals=signals,
            method="deterministic_heuristic",
        )

    # Fallback to uncertain (fail-open: never reject on weak confidence)
    return RelevanceCheckResult(
        decision="uncertain",
        is_useful=True,
        confidence=0.5,
        relevance_score=0.5,
        reason=f"{reason_prefix}: ambiguous or weak domain signals; proceeding with extraction to avoid false rejection.",
        evidence=[],
        signals=signals,
        method="deterministic_heuristic",
    )


def _heuristic_relevance(snippet: str) -> RelevanceCheckResult:
    """Backward-compatible alias for deterministic relevance analysis."""
    return _conservative_deterministic_relevance(snippet, reason_prefix="Heuristic evaluation")


async def _run_relevance_check(text: str) -> RelevanceCheckResult:
    sample = _sample_representative_text(text, max_chars=RELEVANCE_MAX_TOTAL_CHARS)
    llm = get_llm()
    if llm is None:
        return _conservative_deterministic_relevance(
            sample,
            reason_prefix="LLM unavailable; conservative deterministic analysis",
        )

    try:
        system_prompt = load_prompt(PromptId.INGEST_RELEVANCE_V1)
        user_prompt = f"Classify this document/transcript sample and return structured output.\n\nSample:\n{sample}"
        raw = await llm.ainvoke([
            ("system", system_prompt),
            ("user", user_prompt),
        ])
        content = getattr(raw, "content", None) or str(raw)
        content = content.strip()
        if content.startswith("```"):
            lines = content.splitlines()
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            content = "\n".join(lines).strip()

        parsed = json.loads(content)
        raw_decision = str(parsed.get("decision", "")).strip().lower()
        if raw_decision not in ("relevant", "uncertain", "irrelevant"):
            # Infer decision from is_useful / relevance_score if omitted
            if parsed.get("is_useful") is False or float(parsed.get("relevance_score", 1.0)) <= 0.2:
                raw_decision = "irrelevant"
            elif float(parsed.get("relevance_score", 0.7)) >= 0.6:
                raw_decision = "relevant"
            else:
                raw_decision = "uncertain"

        # Rejection requires high confidence and unambiguous irrelevance
        if "confidence" in parsed:
            confidence = max(0.0, min(1.0, float(parsed["confidence"])))
        else:
            confidence = 0.9 if raw_decision in ("relevant", "irrelevant") else 0.5
        relevance_score = max(0.0, min(1.0, float(parsed.get("relevance_score", 0.85 if raw_decision == "relevant" else (0.5 if raw_decision == "uncertain" else 0.1)))))

        if raw_decision == "irrelevant":
            if confidence < 0.75:
                # Asymmetric risk: if not high confidence irrelevant, treat as uncertain (fail-open)
                raw_decision = "uncertain"
                is_useful = True
            else:
                is_useful = False
        else:
            is_useful = True

        evidence_list = parsed.get("evidence")
        evidence = [str(e).strip() for e in evidence_list] if isinstance(evidence_list, list) else []
        signals = parsed.get("signals") if isinstance(parsed.get("signals"), dict) else None

        return RelevanceCheckResult(
            decision=raw_decision,
            is_useful=is_useful,
            confidence=confidence,
            relevance_score=relevance_score,
            reason=str(parsed.get("reason", "Structured LLM evaluation.")).strip(),
            evidence=evidence,
            signals=signals,
            method="llm",
        )
    except Exception as e:
        logger.warning(
            "LLM relevance check failed (%s: %s); failing open with conservative deterministic analysis",
            type(e).__name__,
            e,
        )
        res = _conservative_deterministic_relevance(
            sample,
            reason_prefix=f"LLM failure ({type(e).__name__}); fallback analysis",
        )
        res.method = "fail_open_fallback"
        return res
